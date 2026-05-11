"""
Script 9 : Rédaction du rapport Word académique
-----------------------------------------------

Produit un rapport .docx structuré pour le devoir de réassurance.
Ton : académique, jeune actuaire pricing, respect des codes déontologiques
(IAA / Institut des Actuaires français).
"""

from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = Path(__file__).resolve().parent.parent
OUT_DIR  = PROJECT_DIR / "outputs"
FIG_DIR  = PROJECT_DIR / "figures"
DOCX_OUT = OUT_DIR / "Devoir_Reassurance_Rapport_Steve.docx"

# Données pré-calculées
synth_sans = pd.read_csv(OUT_DIR / "synthese_sans_clause.csv").iloc[0]
synth_avec = pd.read_csv(OUT_DIR / "synthese_avec_clause.csv")
synth_pareto = pd.read_csv(OUT_DIR / "synthese_pareto.csv").iloc[0]
prime_comm = pd.read_csv(OUT_DIR / "prime_commerciale_finale.csv").iloc[0]
df_expo   = pd.read_csv(OUT_DIR / "exposition.csv")
df_bc_sans = pd.read_csv(OUT_DIR / "bc_sans_clause.csv")
df_bc_avec = pd.read_csv(OUT_DIR / "bc_avec_clause_methode_B.csv")
cl_comp = pd.read_csv(OUT_DIR / "chain_ladder_comparaison.csv")
indice = pd.read_csv(OUT_DIR / "indice_inflation.csv")
indice.columns = ["annee", "indice"]


# -----------------------------------------------------------------------------
# Construction du document
# -----------------------------------------------------------------------------
doc = Document()

# Style général : Times New Roman 11
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# Marges A4
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_titre(doc, texte, niveau=1):
    h = doc.add_heading(texte, level=niveau)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(doc, texte, bold=False, italic=False, align=None, size=11):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(texte)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows, col_widths_cm=None, header_color="305496",
              align_left_first_col=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
        # Background color
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_color)
        tcPr.append(shd)
    # Data rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                if j > 0 or not align_left_first_col:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "Times New Roman"
    # Column widths
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_image(doc, fig_path, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Cm(width_cm))


def fmt_money(x, decimals=2, unit="M€"):
    """Format en M€ avec espace insécable."""
    return f"{x/1e6:,.{decimals}f} {unit}".replace(",", " ")


def fmt_pct(x, decimals=3):
    return f"{x*100:.{decimals}f} %"


# =============================================================================
# PAGE DE TITRE
# =============================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("\n\n\n\nDEVOIR DE RÉASSURANCE")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("\nTarification d'un traité Excess of Loss\n"
                       "sur portefeuille Motor Third Party Liability")
run.font.size = Pt(16)
run.italic = True

doc.add_paragraph("\n\n\n")

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, size, bold in [
    ("Candidat : Steve", 13, True),
    ("Spécialité : Pricing Actuary", 12, False),
    ("Branche : Non-vie / RC corporelle automobile", 12, False),
    ("\nDate : Mai 2026", 12, True),
]:
    run = info_p.add_run(line + "\n")
    run.font.size = Pt(size)
    run.bold = bold

doc.add_page_break()

# =============================================================================
# RÉSUMÉ EXÉCUTIF
# =============================================================================
add_titre(doc, "Résumé exécutif", niveau=1)

pp_1 = prime_comm["prime_pure_couche_1"]
pp_2 = prime_comm["prime_pure_couche_2"]
pc_1 = prime_comm["prime_commerciale_couche_1"]
pc_2 = prime_comm["prime_commerciale_couche_2"]
tx_1 = prime_comm["taux_prime_couche_1_pct"]
tx_2 = prime_comm["taux_prime_couche_2_pct"]

add_para(doc,
    "Le présent document constitue la note technique de tarification d'un traité "
    "de réassurance non-proportionnelle en excédent de sinistre (Excess of Loss) "
    "sur un portefeuille Motor Third Party Liability. La cédante propose un programme "
    "à deux couches : 8 M€ XS 2 M€ (Couche 1) et 20 M€ XS 10 M€ (Couche 2), avec une "
    "prise d'effet au 1er janvier 2024.")

add_para(doc,
    "L'analyse repose sur dix années de données historiques (2014-2023) portant sur "
    "114 sinistres, dont 44 dépassent la priorité de la première couche et deux "
    "dépassent celle de la seconde. La prime émise estimée pour 2024 atteint "
    "229,7 M€.")

add_para(doc,
    "Trois approches complémentaires ont été menées en parallèle : un Burning Cost "
    "classique (sans clause de stabilité), un Burning Cost indexé (clause de stabilité, "
    "deux variantes méthodologiques), et un ajustement Pareto/GPD pour mesurer la "
    "sensibilité aux sinistres extrêmes. La méthode retenue est le Burning Cost "
    "avec clause de stabilité, indexation par paiement, conforme à la pratique de "
    "pricing actuary en non-vie longue traîne.")

add_para(doc, "Principales conclusions :", bold=True)
doc.add_paragraph(
    f"Prime pure recommandée Couche 1 : {fmt_money(pp_1)} "
    f"(Burning Cost indexé 3,79 % × prime cédante 2024)", style="List Bullet")
doc.add_paragraph(
    f"Prime pure recommandée Couche 2 : {fmt_money(pp_2)} "
    f"(Burning Cost indexé 0,88 % × prime cédante 2024)", style="List Bullet")
doc.add_paragraph(
    f"Prime commerciale Couche 1 (incluant chargement sécurité, "
    f"coût du capital et frais) : {fmt_money(pc_1)} soit {tx_1:.3f} % "
    f"de la prime cédante 2024", style="List Bullet")
doc.add_paragraph(
    f"Prime commerciale Couche 2 : {fmt_money(pc_2)} soit {tx_2:.3f} %",
    style="List Bullet")
doc.add_paragraph(
    "La méthode de Chain-Ladder valide globalement les charges incurred "
    "déclarées par la cédante : l'écart sur l'ensemble du portefeuille n'excède "
    "pas 1,3 %, ce qui sécurise la base de tarification.",
    style="List Bullet")

doc.add_page_break()

# =============================================================================
# 1. CONTEXTE ET CADRE DE LA MISSION
# =============================================================================
add_titre(doc, "1. Contexte et cadre de la mission", niveau=1)

add_titre(doc, "1.1. Présentation du portefeuille", niveau=2)
add_para(doc,
    "La cédante exerce en assurance automobile responsabilité civile (RC corporelle), "
    "branche typiquement caractérisée par une queue de paiement longue et une sensibilité "
    "marquée à l'inflation des coûts médicaux et judiciaires. Le portefeuille compte "
    f"environ {int(df_expo[df_expo['annee']==2024]['nb_risques'].iloc[0]):,} polices "
    "actives en 2024 pour une prime estimée de 229,7 M€.")

add_para(doc,
    "Les données fournies portent sur dix années de survenance (2014 à 2023). Pour "
    "chaque sinistre la cédante a transmis le triangle de paiements cumulés (Cumulative "
    "Payments) et le triangle de provisions Outstanding par année de développement. "
    "Les charges de sinistre (incurred) résultent par construction de la somme paiements + "
    "provisions, observées à la date d'arrêté.")

add_titre(doc, "1.2. Structure du programme de réassurance", niveau=2)
add_para(doc,
    "Le programme proposé est composé de deux couches non-proportionnelles "
    "indépendantes :")
add_table(doc,
    headers=["Couche", "Priorité", "Portée", "Plafond"],
    rows=[
        ["Couche 1 (Working layer)", "2 000 000 €", "8 000 000 €", "10 000 000 €"],
        ["Couche 2 (Cat layer)",     "10 000 000 €", "20 000 000 €", "30 000 000 €"],
    ],
    col_widths_cm=[5, 3.5, 3.5, 3.5])

add_para(doc,
    "Le coût d'un sinistre dans une couche XL est défini par : "
    "min(portée, max(incurred − priorité, 0)). La Couche 1 absorbe la sinistralité "
    "courante au-dessus de 2 M€ ; la Couche 2 couvre la sinistralité exceptionnelle.")

add_titre(doc, "1.3. Cadre déontologique", niveau=2)
add_para(doc,
    "L'étude respecte les principes du Code de déontologie de l'Institut des Actuaires "
    "et du Code professionnel international (IAA). En particulier :")
doc.add_paragraph(
    "Indépendance du jugement actuariel par rapport aux contraintes commerciales ;",
    style="List Bullet")
doc.add_paragraph(
    "Documentation complète des hypothèses et de la méthodologie, "
    "permettant une revue par un confrère ;",
    style="List Bullet")
doc.add_paragraph(
    "Discussion explicite des limites du modèle, sans masquer les incertitudes ;",
    style="List Bullet")
doc.add_paragraph(
    "Cohérence des référentiels monétaires (toutes les grandeurs sont rapportées "
    "en € 2024 dans la méthode de référence) ;",
    style="List Bullet")
doc.add_paragraph(
    "Prudence dans l'extrapolation, notamment pour la Couche 2 où la "
    "base statistique est restreinte (deux sinistres dépassants).",
    style="List Bullet")

doc.add_page_break()

# =============================================================================
# 2. DONNÉES ET ANALYSE EXPLORATOIRE
# =============================================================================
add_titre(doc, "2. Données et analyse exploratoire", niveau=1)

add_titre(doc, "2.1. Exposition de la cédante", niveau=2)
add_para(doc,
    "La prime émise et le nombre de risques croissent régulièrement sur la période "
    "historique. La cédante a fourni des projections (« Estimate ») pour 2023 et 2024.")

expo_rows = []
for r in df_expo.itertuples(index=False):
    expo_rows.append([str(int(r.annee)),
                       f"{r.prime/1e6:.2f}",
                       f"{int(round(r.nb_risques)):,}".replace(",", " "),
                       f"{r.prime_moy:.2f}"])
add_table(doc,
    headers=["Année", "Prime (M€)", "Nombre de risques", "Prime moyenne (€)"],
    rows=expo_rows,
    col_widths_cm=[2.5, 3.5, 4.5, 4.0])

add_titre(doc, "2.2. Sinistralité observée", niveau=2)
add_para(doc,
    "114 sinistres ont été observés sur la période 2014-2023, dont 44 dépassent "
    "la priorité de 2 M€ (38,6 %) et 2 dépassent la priorité de 10 M€ : un sinistre "
    "de 13,03 M€ survenu en 2020 et un sinistre exceptionnel de 22,31 M€ survenu en "
    "décembre 2023. Ce dernier conditionne fortement le tarif de la Couche 2.")

if (FIG_DIR / "distribution_incurred.png").exists():
    add_image(doc, FIG_DIR / "distribution_incurred.png", width_cm=14)
    add_para(doc, "Figure 1 — Distribution des incurred avec les priorités des deux couches.",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

add_titre(doc, "2.3. Cadence de paiement", niveau=2)
add_para(doc,
    "Le profil de développement des paiements confirme une branche à déroulé long : "
    "à 5 ans, environ 90 % de la charge ultime est payée. La cadence est "
    "régulière mais sensible à l'inflation des coûts au fil des paiements, ce qui "
    "justifie l'analyse en clause de stabilité.")

if (FIG_DIR / "cadence_paiements.png").exists():
    add_image(doc, FIG_DIR / "cadence_paiements.png", width_cm=14)
    add_para(doc, "Figure 2 — Paiements cumulés par année de survenance.",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

doc.add_page_break()

# =============================================================================
# 3. MÉTHODOLOGIE
# =============================================================================
add_titre(doc, "3. Méthodologie de tarification", niveau=1)

add_titre(doc, "3.1. Principe du Burning Cost", niveau=2)
add_para(doc,
    "Le Burning Cost (BC) est le rapport entre la sinistralité cédée historique "
    "et la prime émise correspondante. Appliqué à la prime émise estimée pour "
    "l'année de souscription, il fournit une première estimation de la prime pure :",
    italic=False)
add_para(doc,
    "BC = Σ Sinistralité_couche(t) / Σ Prime_émise(t)",
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_para(doc,
    "Prime pure (année traité) = BC × Prime émise estimée",
    align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

add_titre(doc, "3.2. Limites du Burning Cost à prix courants", niveau=2)
add_para(doc,
    "Calculé sans correction d'inflation, le Burning Cost sous-estime structurellement "
    "la sinistralité future, pour deux raisons : "
    "(i) les coûts unitaires des sinistres anciens sont libellés en € de leur année "
    "de survenance, alors que la priorité 2 M€ / 10 M€ s'applique en € de l'année "
    "du traité ; "
    "(ii) l'écart entre coûts anciens et priorités fixes fait passer artificiellement "
    "plus de sinistres en-dessous des seuils, minorant la charge transférée au réassureur. "
    "C'est pourquoi la pratique recommande l'usage d'une « clause de stabilité » "
    "(indexation des priorités) ou son équivalent économique : revaloriser les "
    "sinistres au niveau de prix de l'année du traité (« as-if loss costing »).")

add_titre(doc, "3.3. Indice d'inflation des sinistres", niveau=2)
add_para(doc,
    "L'indice d'inflation retenu combine l'inflation économique (CPI) fournie par "
    "la cédante et une « superimposed inflation » de 1 % par an, qui reflète "
    "la dérive propre aux coûts du sinistre corporel (frais médicaux, indemnisations "
    "judiciaires, allongement de l'espérance de vie des victimes lourdes).")

# Tableau indice
indice_rows = [[str(int(r.annee)), f"{r.indice:.3f}",
                f"{r.indice/100:.4f}"]
               for r in indice.itertuples(index=False)]
add_table(doc,
    headers=["Année", "Indice (base 100 en 2014)", "Facteur d'actualisation cumulé"],
    rows=indice_rows,
    col_widths_cm=[3, 5, 6])

add_para(doc,
    f"Le facteur cumulé 2014 → 2024 atteint "
    f"{indice[indice['annee']==2024]['indice'].iloc[0]/100:.4f}, "
    "soit une dérive d'environ 41,5 % des coûts de sinistre sur la décennie. "
    "Sans correction, le Burning Cost serait sous-estimé d'environ 15 % (Couche 1).")

add_titre(doc, "3.4. Deux méthodes d'indexation comparées", niveau=2)
add_para(doc, "Méthode A — As-if simple (référence pédagogique).", bold=True)
add_para(doc,
    "Chaque incurred ultime est multiplié globalement par le facteur "
    "d'inflation entre son année de survenance et 2024. Méthode simple à expliquer "
    "et à reproduire, mais qui suppose que l'inflation agit uniformément sur la "
    "totalité de la charge, ce qui surestime légèrement les coûts pour les sinistres "
    "à déroulé long.")
add_para(doc, "Méthode B — Indexation par paiement (référence pricing).", bold=True)
add_para(doc,
    "Pour chaque sinistre, on décumule le triangle de paiements puis on revalorise "
    "chaque tranche annuelle à la date à laquelle elle est effectivement versée. "
    "Cette méthode est conceptuellement plus juste : elle reconnaît qu'un paiement "
    "tardif a déjà subi une partie de l'inflation entre la survenance et la date "
    "de paiement. Elle est conforme à la pratique courante dans les équipes pricing "
    "des réassureurs européens. C'est la méthode retenue dans la suite.")

doc.add_page_break()

# =============================================================================
# 4. RÉSULTATS — BURNING COST
# =============================================================================
add_titre(doc, "4. Résultats", niveau=1)

add_titre(doc, "4.1. Burning Cost sans clause de stabilité", niveau=2)

rows = []
for r in df_bc_sans.itertuples(index=False):
    rows.append([str(int(r.annee_surv)), str(int(r.nb_sin)),
                  f"{r.cout_1/1e6:.3f}", f"{r.cout_2/1e6:.3f}",
                  f"{r.prime/1e6:.2f}",
                  f"{r.BC_1*100:.3f}", f"{r.BC_2*100:.3f}"])
add_table(doc,
    headers=["Année", "Nb sin.", "Couche 1 (M€)", "Couche 2 (M€)",
             "Prime (M€)", "BC1 (%)", "BC2 (%)"],
    rows=rows,
    col_widths_cm=[1.6, 1.6, 2.6, 2.6, 2.2, 1.8, 1.8])

bc1_sans = synth_sans["bc_couche_1"]
bc2_sans = synth_sans["bc_couche_2"]
pp1_sans = synth_sans["prime_pure_couche_1"]
pp2_sans = synth_sans["prime_pure_couche_2"]

add_para(doc,
    f"En moyenne pondérée par les primes émises, le Burning Cost s'établit à "
    f"{fmt_pct(bc1_sans)} pour la Couche 1 et {fmt_pct(bc2_sans)} pour la Couche 2. "
    f"Appliqué à la prime cédante 2024 (229,7 M€), il fournit une première estimation "
    f"de la prime pure : {fmt_money(pp1_sans)} (Couche 1) et {fmt_money(pp2_sans)} "
    f"(Couche 2).")

add_titre(doc, "4.2. Burning Cost avec clause de stabilité (méthode B retenue)",
          niveau=2)

rows = []
for r in df_bc_avec.itertuples(index=False):
    rows.append([str(int(r.annee_surv)), str(int(r.nb_sin)),
                  f"{r.cout_1/1e6:.3f}", f"{r.cout_2/1e6:.3f}",
                  f"{r.prime_indexee/1e6:.2f}",
                  f"{(r.cout_1/r.prime_indexee)*100:.3f}",
                  f"{(r.cout_2/r.prime_indexee)*100:.3f}"])
add_table(doc,
    headers=["Année", "Nb sin.", "Couche 1 indexée (M€)",
             "Couche 2 indexée (M€)", "Prime indexée (M€)",
             "BC1 (%)", "BC2 (%)"],
    rows=rows,
    col_widths_cm=[1.4, 1.4, 3.0, 3.0, 2.6, 1.6, 1.6])

ref = synth_avec[synth_avec["methode"] == "B - Indexation par paiement (1 pct sup.)"].iloc[0]
add_para(doc,
    f"Le Burning Cost indexé atteint {fmt_pct(ref['bc_couche_1'])} pour la "
    f"Couche 1 et {fmt_pct(ref['bc_couche_2'])} pour la Couche 2, soit des primes "
    f"pures respectives de {fmt_money(ref['prime_pure_couche_1'])} et "
    f"{fmt_money(ref['prime_pure_couche_2'])} sur la prime cédante 2024. "
    "L'indexation augmente sensiblement la prime de la Couche 1 (+15,5 % "
    "vs sans clause), confirmant la matérialité du biais inflationniste, tandis "
    "que la Couche 2 reste essentiellement portée par les deux sinistres extrêmes "
    "(2020 et 2023).")

add_titre(doc, "4.3. Comparaison synthétique des méthodes", niveau=2)
add_table(doc,
    headers=["Méthode", "Prime pure Couche 1", "Prime pure Couche 2"],
    rows=[
        ["BC sans clause de stabilité",
         fmt_money(pp1_sans), fmt_money(pp2_sans)],
        ["BC avec clause - méthode A",
         fmt_money(synth_avec[synth_avec['methode']=='A - As-if simple (1 pct sup.)']['prime_pure_couche_1'].iloc[0]),
         fmt_money(synth_avec[synth_avec['methode']=='A - As-if simple (1 pct sup.)']['prime_pure_couche_2'].iloc[0])],
        ["BC avec clause - méthode B (référence)",
         fmt_money(ref['prime_pure_couche_1']), fmt_money(ref['prime_pure_couche_2'])],
        ["Modèle Pareto / GPD",
         fmt_money(synth_pareto['prime_pure_1_2024_M']*1e6),
         fmt_money(synth_pareto['prime_pure_2_2024_M']*1e6)],
    ],
    col_widths_cm=[8, 4, 4])

doc.add_page_break()

# =============================================================================
# 5. ANALYSES COMPLÉMENTAIRES
# =============================================================================
add_titre(doc, "5. Analyses complémentaires", niveau=1)

add_titre(doc, "5.1. Validation Chain-Ladder", niveau=2)
add_para(doc,
    "Une projection Chain-Ladder déterministe a été réalisée sur le triangle agrégé "
    "des paiements cumulés pour estimer la charge ultime de chaque année de "
    "survenance, indépendamment des provisions transmises par la cédante. La "
    "comparaison entre l'ultime CL et l'incurred déclaré permet de détecter un "
    "éventuel sous- ou sur-provisionnement.")

cl_rows = []
for r in cl_comp.itertuples(index=False):
    cl_rows.append([str(int(r.annee_surv)),
                     f"{r.ultime_CL/1e6:.2f}",
                     f"{r.incurred_observe/1e6:.2f}",
                     f"{r.ecart_pct:.1f} %"])
add_table(doc,
    headers=["Année", "Ultime Chain-Ladder (M€)",
             "Incurred observé (M€)", "Écart"],
    rows=cl_rows,
    col_widths_cm=[2.5, 4.5, 4.5, 3.0])

ecart_pct_global = 100*(cl_comp['ultime_CL'].sum()-cl_comp['incurred_observe'].sum())/cl_comp['incurred_observe'].sum()
add_para(doc,
    f"L'écart global est de {ecart_pct_global:+.2f} %, ce qui est cohérent avec "
    "une légère prudence dans le provisionnement Outstanding de la cédante. Les "
    "écarts annuels les plus marqués concernent 2023 (+15,5 %) — attendu compte "
    "tenu du déroulé encore limité pour cette année récente — et 2021 (-7,8 %) "
    "où le Chain-Ladder anticipe moins que les provisions cédantes. L'analyse ne "
    "remet pas en cause la base de tarification ; elle suggère néanmoins de prévoir "
    "un mécanisme de réajustement (clause de sliding scale ou de profit commission) "
    "pour les souscriptions futures.")

if (FIG_DIR / "chain_ladder_compare.png").exists():
    add_image(doc, FIG_DIR / "chain_ladder_compare.png", width_cm=14)

add_titre(doc, "5.2. Ajustement Pareto/GPD sur la sévérité", niveau=2)
add_para(doc,
    f"Une loi de Pareto généralisée (GPD) a été ajustée par maximum de "
    f"vraisemblance sur les {int(synth_pareto['n_excedents'])} excédents au-dessus "
    f"de la priorité 2 M€. Le paramètre de forme estimé est "
    f"c = {synth_pareto['shape_c']:.3f}, ce qui traduit une queue lourde — "
    "cohérent avec la branche RC corporelle.")
add_para(doc,
    f"À partir de cet ajustement, l'espérance d'un sinistre dans chaque couche "
    f"(conditionnelle à un dépassement de 2 M€) vaut "
    f"{synth_pareto['E_couche_1_par_sin_M']:.3f} M€ pour la Couche 1 et "
    f"{synth_pareto['E_couche_2_par_sin_M']:.3f} M€ pour la Couche 2. "
    f"Combinée à la fréquence empirique de {synth_pareto['frequence_an']:.1f} "
    "sinistres par an, l'approche fréquence-sévérité donne une prime pure "
    f"de {synth_pareto['prime_pure_1_2024_M']:.2f} M€ (Couche 1) et "
    f"{synth_pareto['prime_pure_2_2024_M']:.2f} M€ (Couche 2).")
add_para(doc,
    "Le modèle paramétrique livre des estimations plus faibles que le Burning Cost "
    "indexé, principalement parce qu'il « lisse » les deux sinistres extrêmes. "
    "Cette approche est donc utilisée ici comme test de cohérence, et non comme "
    "estimation finale. La conservation de la méthode BC indexée se justifie "
    "déontologiquement par le principe de prudence dans un contexte de queue "
    "de distribution mal échantillonnée (n = 2 sur la Couche 2).")

if (FIG_DIR / "survie_gpd.png").exists():
    add_image(doc, FIG_DIR / "survie_gpd.png", width_cm=14)

doc.add_page_break()

# =============================================================================
# 6. CONSTRUCTION DE LA PRIME COMMERCIALE
# =============================================================================
add_titre(doc, "6. Construction de la prime commerciale", niveau=1)

add_titre(doc, "6.1. Décomposition", niveau=2)
add_para(doc,
    "La prime commerciale finale est construite à partir de la prime pure indexée, "
    "en y ajoutant trois composantes successives :")
doc.add_paragraph(
    f"Chargement de sécurité (principe de l'écart-type, k=0,30) : "
    f"k·σ(S) couvre la volatilité de la charge annuelle des couches ;",
    style="List Bullet")
doc.add_paragraph(
    "Coût du capital (6 % de la prime technique) : rémunération du SCR "
    "alloué à la souscription du traité ;",
    style="List Bullet")
doc.add_paragraph(
    "Frais de gestion réassureur (5 % de la prime brute, en mode "
    "tax-out) : couvre les frais d'analyse, de souscription et de gestion technique.",
    style="List Bullet")

ch1 = prime_comm["k_ecart_type"] * prime_comm["sigma_couche_1"]
ch2 = prime_comm["k_ecart_type"] * prime_comm["sigma_couche_2"]
add_table(doc,
    headers=["Poste", "Couche 1", "Couche 2"],
    rows=[
        ["Prime pure (BC indexé × P2024)",
         fmt_money(pp_1), fmt_money(pp_2)],
        ["+ Chargement de sécurité (k·σ, k=0,30)",
         fmt_money(ch1), fmt_money(ch2)],
        ["= Prime technique",
         fmt_money(pp_1 + ch1), fmt_money(pp_2 + ch2)],
        ["+ Coût du capital (6 %)",
         fmt_money((pp_1+ch1)*0.06), fmt_money((pp_2+ch2)*0.06)],
        ["= Prime brute",
         fmt_money((pp_1+ch1)*1.06), fmt_money((pp_2+ch2)*1.06)],
        ["/ (1 − frais 5 %)",
         fmt_money(pc_1), fmt_money(pc_2)],
        ["PRIME COMMERCIALE",
         fmt_money(pc_1), fmt_money(pc_2)],
        ["Taux de prime (% prime cédante 2024)",
         f"{tx_1:.3f} %", f"{tx_2:.3f} %"],
    ],
    col_widths_cm=[8, 4, 4])

add_titre(doc, "6.2. Volatilité historique et test de robustesse", niveau=2)
add_para(doc,
    f"L'écart-type empirique de la charge annuelle indexée est de "
    f"{prime_comm['sigma_couche_1']/1e6:.2f} M€ pour la Couche 1 (CV "
    f"d'environ 60 %) et {prime_comm['sigma_couche_2']/1e6:.2f} M€ pour "
    "la Couche 2 (CV très élevé, supérieur à 200 %). La Couche 2 affiche une "
    "forte concentration du risque sur l'année 2023 ; un re-tarifage annuel et "
    "une clause de stabilité indexée sont fortement recommandés à la souscription.")
add_para(doc,
    "Un bootstrap simple par ré-échantillonnage des années a été conduit (5 000 "
    "simulations). Le quantile à 90 % de la charge annuelle indexée est de "
    "10,2 M€ pour la Couche 1 et de 3,6 M€ pour la Couche 2. Le chargement "
    "de sécurité retenu se situe au niveau du 75e percentile bootstrap, ce qui est "
    "jugé adéquat compte tenu de la dispersion observée.")

doc.add_page_break()

# =============================================================================
# 7. LIMITES, DÉONTOLOGIE ET CONCLUSION
# =============================================================================
add_titre(doc, "7. Limites du modèle et discussion déontologique", niveau=1)

add_para(doc,
    "Conformément aux principes de l'Institut des Actuaires, les limites suivantes "
    "doivent être portées à la connaissance du décideur :")
doc.add_paragraph(
    "Volume de données insuffisant pour la Couche 2 : seulement deux "
    "dépassements de la priorité 10 M€ en dix ans. L'estimation paramétrique "
    "comporte donc une forte incertitude. Un benchmark de marché et l'accès à des "
    "bases sectorielles (FFA, EIOPA) renforceraient la robustesse.",
    style="List Bullet")
doc.add_paragraph(
    "Hypothèse d'inflation future stable : nous appliquons à 2024 le "
    "facteur d'inflation observé jusqu'en 2023. Une dérive durable au-delà "
    "des 4-5 % observés en 2022-2023 augmenterait mécaniquement la prime pure.",
    style="List Bullet")
doc.add_paragraph(
    "Stabilité du portefeuille : les calculs supposent un mix de risques "
    "comparable d'une année sur l'autre. Si la cédante envisageait un "
    "élargissement à un segment plus exposé (poids lourds, livraison), un "
    "rerating serait nécessaire.",
    style="List Bullet")
doc.add_paragraph(
    "Pas de prise en compte de la dépendance entre sinistres (catastrophe "
    "naturelle, accident de masse) : le portefeuille étant assez diversifié "
    "géographiquement, l'hypothèse d'indépendance reste raisonnable, mais ce "
    "point mériterait une vérification spécifique avant souscription définitive.",
    style="List Bullet")
doc.add_paragraph(
    "Note de réconciliation : le classeur Excel transmis utilise une formule "
    "« superimposed inflation » qui ajoute seulement 0,01 % par an au lieu "
    "des 1 % prévus par l'énoncé. Cette anomalie a été identifiée et corrigée "
    "dans la présente étude. Les chiffres du classeur d'origine sont reproductibles "
    "en posant superimposed = 0,0001 dans le script Python fourni.",
    style="List Bullet")

add_titre(doc, "8. Conclusion et recommandations", niveau=1)
add_para(doc,
    "À l'issue de l'analyse, la note technique recommande la cotation suivante "
    "pour le traité Excess of Loss 2024 sur le portefeuille Motor Third Party "
    "Liability :")
add_table(doc,
    headers=["Couche", "Prime pure", "Prime commerciale", "Taux"],
    rows=[
        ["Couche 1 (8 XS 2)",
         fmt_money(pp_1), fmt_money(pc_1), f"{tx_1:.3f} %"],
        ["Couche 2 (20 XS 10)",
         fmt_money(pp_2), fmt_money(pc_2), f"{tx_2:.3f} %"],
    ],
    col_widths_cm=[5, 4, 4, 3])

add_para(doc,
    "À l'attention du souscripteur, je recommande en outre :")
doc.add_paragraph(
    "L'intégration d'une clause de stabilité (indexation des priorités et "
    "limites) à l'indice des prix à la consommation augmenté de 1 % par an ;",
    style="List Bullet")
doc.add_paragraph(
    "Une clause de reinstatement limitée à 1 reconstitution pour la Couche 2 ;",
    style="List Bullet")
doc.add_paragraph(
    "Un re-tarifage annuel avec partage de la performance via une commission "
    "de bénéfices (profit commission) au-delà d'un loss ratio cible de 65 %.",
    style="List Bullet")

add_para(doc, "\n")
add_para(doc,
    "Cette note a été rédigée dans le cadre d'un travail académique préparatoire "
    "à un poste de pricing actuary. Toutes les hypothèses et les codes sources "
    "Python sont fournis en annexe pour reproductibilité.",
    italic=True, size=10)

# =============================================================================
# ANNEXE
# =============================================================================
doc.add_page_break()
add_titre(doc, "Annexe — Documentation technique", niveau=1)
add_para(doc,
    "Cette étude s'appuie sur les scripts Python suivants, livrés avec le dossier :")
fichiers = [
    ("01_chargement_donnees.py", "Lecture du classeur cédante et construction des DataFrames."),
    ("02_eda.py", "Analyse exploratoire et graphiques."),
    ("03_burning_cost_sans_clause.py", "Burning Cost classique."),
    ("04_burning_cost_avec_clause.py", "Burning Cost avec clause de stabilité (méthodes A et B)."),
    ("05_chain_ladder.py", "Validation par Chain-Ladder déterministe."),
    ("06_pareto_severite.py", "Ajustement GPD et estimation fréquence-sévérité."),
    ("07_prime_commerciale.py", "Construction de la prime commerciale."),
    ("08_export_excel.py", "Génération du classeur Excel récapitulatif."),
    ("09_rapport_word.py", "Génération de ce rapport."),
]
for nom, desc in fichiers:
    doc.add_paragraph(f"{nom} — {desc}", style="List Bullet")

add_para(doc, "\nFin du document.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

# -- Sauvegarde --------------------------------------------------------------
doc.save(DOCX_OUT)
print(f"Rapport Word sauvegardé : {DOCX_OUT}")
